"""
Render a SelectedResume to PDF (WeasyPrint) or DOCX (python-docx fallback).

Cache strategy: hash selected_resume JSON → check resume_versions table →
if render_hash found with existing file: return cached path.
"""
from __future__ import annotations

import hashlib
import json
import os
import uuid
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.models import ResumeVersion
from backend.schemas import SelectedResume
from backend.utils.logging import get_logger

logger = get_logger(__name__)

RESUMES_DIR = Path("data/resumes")
TEMPLATES_DIR = Path("templates")


def _render_hash(selected_resume: SelectedResume) -> str:
    raw = json.dumps(selected_resume.model_dump(), sort_keys=True)
    return hashlib.sha256(raw.encode()).hexdigest()


def _render_html(selected_resume: SelectedResume) -> str:
    """Render the Jinja2 HTML template with resume data."""
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES_DIR)),
        autoescape=True,
    )
    template = env.get_template("resume.html")
    return template.render(**selected_resume.model_dump())


def _html_to_pdf(html: str, out_path: Path) -> bool:
    """Render HTML to PDF via WeasyPrint. Returns True on success."""
    try:
        from weasyprint import HTML
        HTML(string=html, base_url=str(TEMPLATES_DIR)).write_pdf(str(out_path))
        return True
    except Exception as exc:
        logger.warning("weasyprint_failed", extra={"error": str(exc)})
        return False


def _html_to_docx(selected_resume: SelectedResume, out_path: Path) -> bool:
    """Basic python-docx fallback. Returns True on success."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        for sec in doc.sections:
            sec.left_margin = Inches(0.6)
            sec.right_margin = Inches(0.6)
            sec.top_margin = Inches(0.5)
            sec.bottom_margin = Inches(0.5)

        personal = selected_resume.personal
        name_para = doc.add_paragraph()
        run = name_para.add_run(personal.get("name", ""))
        run.bold = True
        run.font.size = Pt(16)
        name_para.alignment = 1  # center

        contacts = [personal.get("email", ""), personal.get("phone", ""), personal.get("location", "")]
        contact_para = doc.add_paragraph(" | ".join(c for c in contacts if c))
        contact_para.alignment = 1

        for section_name in selected_resume.section_order:
            if section_name == "work_experience":
                doc.add_heading("Experience", level=2)
                for exp in selected_resume.work_experience:
                    p = doc.add_paragraph()
                    p.add_run(f"{exp.get('company')} — {exp.get('role')}").bold = True
                    for bullet in exp.get("bullets", []):
                        doc.add_paragraph(bullet["text"], style="List Bullet")
            elif section_name == "projects":
                doc.add_heading("Projects", level=2)
                for proj in selected_resume.projects:
                    p = doc.add_paragraph()
                    p.add_run(proj.get("name", "")).bold = True
                    for bullet in proj.get("bullets", []):
                        doc.add_paragraph(bullet["text"], style="List Bullet")
            elif section_name == "education":
                doc.add_heading("Education", level=2)
                for edu in selected_resume.education:
                    p = doc.add_paragraph()
                    p.add_run(f"{edu.get('institution')} — {edu.get('degree')}").bold = True
            elif section_name == "skills" and selected_resume.skills:
                doc.add_heading("Skills", level=2)
                skills = selected_resume.skills
                for cat, items in skills.items():
                    if items and isinstance(items, list):
                        doc.add_paragraph(f"{cat.capitalize()}: {', '.join(items)}")

        doc.save(str(out_path))
        return True
    except Exception as exc:
        logger.warning("docx_render_failed", extra={"error": str(exc)})
        return False


async def render_resume(
    selected_resume: SelectedResume,
    application_id: str,
    session: AsyncSession,
    fmt: str = "pdf",
) -> str:
    """
    Render selected_resume to PDF or DOCX.

    1. Compute render_hash → check resume_versions cache.
    2. Cache hit with existing file → return path.
    3. Render HTML via Jinja2 → WeasyPrint PDF → DOCX fallback.
    4. Store in resume_versions → return absolute path.
    """
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    render_hash = _render_hash(selected_resume)

    # --- Cache check ---
    result = await session.execute(
        select(ResumeVersion).where(ResumeVersion.render_hash == render_hash)
    )
    cached = result.scalar_one_or_none()
    if cached:
        if fmt == "pdf" and cached.pdf_path and os.path.exists(cached.pdf_path):
            logger.info("render_cache_hit", extra={"application_id": application_id})
            return cached.pdf_path
        if fmt == "docx" and cached.docx_path and os.path.exists(cached.docx_path):
            return cached.docx_path

    html = _render_html(selected_resume)
    pdf_path: str | None = None
    docx_path: str | None = None

    if fmt == "pdf":
        out_pdf = RESUMES_DIR / f"{application_id}.pdf"
        if _html_to_pdf(html, out_pdf):
            pdf_path = str(out_pdf)
        else:
            out_docx = RESUMES_DIR / f"{application_id}.docx"
            if _html_to_docx(selected_resume, out_docx):
                docx_path = str(out_docx)
    else:
        out_docx = RESUMES_DIR / f"{application_id}.docx"
        if _html_to_docx(selected_resume, out_docx):
            docx_path = str(out_docx)

    final_path = pdf_path or docx_path
    if not final_path:
        raise RuntimeError(f"Resume render failed for application {application_id}")

    session.add(
        ResumeVersion(
            id=str(uuid.uuid4()),
            application_id=application_id,
            tailored_json=json.dumps(selected_resume.model_dump()),
            pdf_path=pdf_path,
            docx_path=docx_path,
            render_hash=render_hash,
        )
    )
    await session.commit()

    logger.info(
        "resume_rendered",
        extra={"application_id": application_id, "fmt": fmt, "path": final_path},
    )
    return final_path
